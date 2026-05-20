"""Tests for SectionProperties page_margins in the declarative API."""

import pytest
from docx.shared import Inches

from cmi_docx import declarative


@pytest.mark.asyncio
async def test_all_margins() -> None:
    """Test that all four margins are applied when all keys are set.

    When page_margins specifies all four sides to Inches(2), every margin
    attribute on the resulting section should equal Inches(2).
    """
    doc = declarative.Document(
        sections=[
            declarative.Section(
                children=[declarative.Paragraph(text="All margins")],
                properties=declarative.SectionProperties(
                    page_margins={
                        "top": Inches(2),
                        "bottom": Inches(2),
                        "left": Inches(2),
                        "right": Inches(2),
                    },
                ),
            ),
        ],
    )

    docx_doc = await doc.to_docx()
    section = docx_doc.sections[-2]
    assert section.top_margin == Inches(2)
    assert section.bottom_margin == Inches(2)
    assert section.left_margin == Inches(2)
    assert section.right_margin == Inches(2)


@pytest.mark.asyncio
async def test_partial_margins() -> None:
    """Test that only the specified margins are changed when keys are omitted.

    When page_margins specifies only 'top' and 'left' to Inches(1.5), those
    two margins should equal Inches(1.5) while 'bottom' and 'right' retain
    their default values (which differ from Inches(1.5)).
    """
    doc = declarative.Document(
        sections=[
            declarative.Section(
                children=[declarative.Paragraph(text="Partial margins")],
                properties=declarative.SectionProperties(
                    page_margins={
                        "top": Inches(1.5),
                        "left": Inches(1.5),
                    },
                ),
            ),
        ],
    )

    docx_doc = await doc.to_docx()
    section = docx_doc.sections[-2]
    assert section.top_margin == Inches(1.5)
    assert section.left_margin == Inches(1.5)
    assert section.bottom_margin != Inches(1.5)
    assert section.right_margin != Inches(1.5)


@pytest.mark.asyncio
async def test_no_margins() -> None:
    """Test that default margins are preserved when page_margins is not set.

    When no page_margins field is specified on SectionProperties, all four
    margin attributes on the resulting section should be non-None (the docx
    template defaults are preserved).
    """
    doc = declarative.Document(
        sections=[
            declarative.Section(
                children=[declarative.Paragraph(text="Default margins")],
            ),
        ],
    )

    docx_doc = await doc.to_docx()
    section = docx_doc.sections[-2]
    assert section.top_margin is not None
    assert section.bottom_margin is not None
    assert section.left_margin is not None
    assert section.right_margin is not None


@pytest.mark.asyncio
async def test_margins_apply_to_correct_section() -> None:
    """Test that margins are applied to the correct section in a multi-section document.

    With 2 declarative sections, the first has no page_margins and the second
    has top_margin=Inches(3). Only sections[1].top_margin should equal Inches(3);
    sections[0].top_margin should retain its default value.
    """
    doc = declarative.Document(
        sections=[
            declarative.Section(
                children=[declarative.Paragraph(text="Section without margins")],
            ),
            declarative.Section(
                children=[declarative.Paragraph(text="Section with top margin")],
                properties=declarative.SectionProperties(
                    page_margins={"top": Inches(3)},
                ),
            ),
        ],
    )

    docx_doc = await doc.to_docx()
    sections = list(docx_doc.sections)

    assert sections[0].top_margin != Inches(3)
    assert sections[1].top_margin == Inches(3)
