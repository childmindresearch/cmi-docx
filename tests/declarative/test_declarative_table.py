"""Table tests for the declarative API."""

import pytest
from docx import shared
from docx.oxml.ns import qn
from docx.oxml.simpletypes import ST_Merge

from cmi_docx import declarative


@pytest.mark.asyncio
async def test_simple_table() -> None:
    """Test creating a document with a simple table."""
    doc = declarative.Document(
        sections=[
            declarative.Section(
                children=[
                    declarative.Table(
                        rows=[
                            declarative.TableRow(
                                children=[
                                    declarative.TableCell(
                                        children=[
                                            declarative.Paragraph(text="Header 1"),
                                            declarative.Paragraph(text="Paragraph 2"),
                                        ]
                                    ),
                                ],
                            ),
                        ],
                    ),
                ],
            ),
        ],
    )

    docx = await doc.to_docx()
    assert docx.tables[0].rows[0].cells[0].paragraphs[0].text == "Header 1"
    assert docx.tables[0].rows[0].cells[0].paragraphs[1].text == "Paragraph 2"


@pytest.mark.asyncio
async def test_table_horizontal_merge() -> None:
    """Test that grid_span is applied correctly to a cell in the rendered table.

    A cell with grid_span=2 followed by one normal cell should produce a table
    with 3 columns total (2 from the span + 1 from the second cell), and the
    first physical TC element should carry grid_span=2 in the XML.
    """
    doc = declarative.Document(
        sections=[
            declarative.Section(
                children=[
                    declarative.Table(
                        rows=[
                            declarative.TableRow(
                                children=[
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="Merged")],
                                        grid_span=2,
                                    ),
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="Single")],
                                    ),
                                ],
                            ),
                        ],
                    ),
                ],
            ),
        ],
    )

    docx_doc = await doc.to_docx()
    tbl = docx_doc.tables[0]

    expected_grid_span = 2
    expected_columns = 3  # grid_span=2 + 1 normal cell

    assert len(tbl.columns) == expected_columns
    assert tbl.cell(0, 0)._tc.grid_span == expected_grid_span

    # The first TC element contains the "Merged" paragraph text (first run).
    # Access via raw XML to get the physical TC rather than a virtual merged cell.
    tc0 = tbl._tbl.findall(qn("w:tr"))[0].findall(qn("w:tc"))[0]
    first_run = tc0.findall(qn("w:p"))[0].findall(qn("w:r"))[0]
    assert first_run.find(qn("w:t")).text == "Merged"

    row_tcs = tbl._tbl.findall(qn("w:tr"))[0].findall(qn("w:tc"))
    # After the fix, TC[1] is the "Single" cell (surplus TC was removed)
    tc1_run = row_tcs[1].findall(qn("w:p"))[0].findall(qn("w:r"))[0]
    assert tc1_run.find(qn("w:t")).text == "Single"


@pytest.mark.asyncio
async def test_table_vertical_merge() -> None:
    """Test that vmerge sets restart and continue on the correct TC elements.

    Row 0 should carry vMerge=restart and row 1 should carry vMerge=continue.
    Both are verified via the raw CT_Tc elements accessed through the XML tree
    (bypassing python-docx's merged-cell resolution in tbl.cell()).
    """
    doc = declarative.Document(
        sections=[
            declarative.Section(
                children=[
                    declarative.Table(
                        rows=[
                            declarative.TableRow(
                                children=[
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="Top")],
                                        vmerge="restart",
                                    ),
                                ],
                            ),
                            declarative.TableRow(
                                children=[
                                    declarative.TableCell(
                                        children=None,
                                        vmerge="continue",
                                    ),
                                ],
                            ),
                        ],
                    ),
                ],
            ),
        ],
    )

    docx_doc = await doc.to_docx()
    tbl = docx_doc.tables[0]

    # Access TC elements directly via the XML tree to avoid python-docx's
    # merged-cell resolution (tbl.cell(row, col) follows vMerge upward).
    xml_rows = tbl._tbl.findall(qn("w:tr"))
    tc_row0 = xml_rows[0].findall(qn("w:tc"))[0]
    tc_row1 = xml_rows[1].findall(qn("w:tc"))[0]

    assert tc_row0.vMerge == ST_Merge.RESTART
    assert tc_row1.vMerge == ST_Merge.CONTINUE


@pytest.mark.asyncio
async def test_table_combined_merge() -> None:
    """Test a table combining horizontal merge in row 0 and normal cells in row 1.

    Row 0 has a cell with grid_span=2 and one normal cell, yielding 3 columns.
    Row 1 has three normal cells. This verifies grid_span and ordinary content
    coexist correctly across rows.
    """
    doc = declarative.Document(
        sections=[
            declarative.Section(
                children=[
                    declarative.Table(
                        rows=[
                            declarative.TableRow(
                                children=[
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="Wide")],
                                        grid_span=2,
                                    ),
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="Right")],
                                    ),
                                ],
                            ),
                            declarative.TableRow(
                                children=[
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="A")],
                                    ),
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="B")],
                                    ),
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="C")],
                                    ),
                                ],
                            ),
                        ],
                    ),
                ],
            ),
        ],
    )

    docx_doc = await doc.to_docx()
    tbl = docx_doc.tables[0]

    expected_grid_span = 2
    expected_columns = 3  # grid_span=2 + 1 normal cell

    assert len(tbl.columns) == expected_columns
    assert tbl.cell(0, 0)._tc.grid_span == expected_grid_span

    # The first TC in row 0 contains the "Wide" paragraph text (first run).
    # Access via raw XML to get the physical TC rather than a virtual merged cell.
    xml_rows = tbl._tbl.findall(qn("w:tr"))
    tc_row0 = xml_rows[0].findall(qn("w:tc"))[0]
    first_run = tc_row0.findall(qn("w:p"))[0].findall(qn("w:r"))[0]
    assert first_run.find(qn("w:t")).text == "Wide"

    row0_tcs = tbl._tbl.findall(qn("w:tr"))[0].findall(qn("w:tc"))
    # After the fix, TC[1] is the "Right" cell (surplus TC was removed)
    right_run = row0_tcs[1].findall(qn("w:p"))[0].findall(qn("w:r"))[0]
    assert right_run.find(qn("w:t")).text == "Right"

    # Row 1 has three independent cells with expected content
    assert tbl.rows[1].cells[0].paragraphs[0].text == "A"
    assert tbl.rows[1].cells[1].paragraphs[0].text == "B"
    assert tbl.rows[1].cells[2].paragraphs[0].text == "C"


@pytest.mark.asyncio
async def test_horizontal_merge_removes_surplus_tc() -> None:
    """Test that a grid_span=2 cell causes the surplus TC to be removed.

    When cell 0 spans 2 grid columns, the row must contain exactly 2 physical
    TC elements (one spanned, one normal) — not 3. A third TC would make Word
    render an extra phantom column.
    """
    doc = declarative.Document(
        sections=[
            declarative.Section(
                children=[
                    declarative.Table(
                        rows=[
                            declarative.TableRow(
                                children=[
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="Merged")],
                                        grid_span=2,
                                    ),
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="Single")],
                                    ),
                                ],
                            ),
                        ],
                    ),
                ],
            ),
        ],
    )

    docx_doc = await doc.to_docx()
    tbl = docx_doc.tables[0]
    row_tcs = tbl._tbl.findall(qn("w:tr"))[0].findall(qn("w:tc"))
    assert len(row_tcs) == 2  # noqa: PLR2004


@pytest.mark.asyncio
async def test_table_column_widths() -> None:
    """Test that column_widths sets autofit=False and each column width in twips.

    A 1-row, 3-column table with column_widths=[1440, 2880, 1440] should disable
    autofit and set each column to the corresponding width in DXA/twip units.
    """
    doc = declarative.Document(
        sections=[
            declarative.Section(
                children=[
                    declarative.Table(
                        rows=[
                            declarative.TableRow(
                                children=[
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="A")],
                                    ),
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="B")],
                                    ),
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="C")],
                                    ),
                                ],
                            ),
                        ],
                        column_widths=[1440, 2880, 1440],
                    ),
                ],
            ),
        ],
    )

    docx_doc = await doc.to_docx()
    docx_table = docx_doc.tables[0]

    assert docx_table.autofit is False
    assert docx_table.columns[0].width == shared.Twips(1440)
    assert docx_table.columns[1].width == shared.Twips(2880)
    assert docx_table.columns[2].width == shared.Twips(1440)


@pytest.mark.asyncio
async def test_table_column_widths_with_grid_span() -> None:
    """Test that column_widths works correctly with a grid_span cell.

    Row 0 has a cell spanning 2 columns and one normal cell (3 logical columns).
    Row 1 has three normal cells. column_widths=[1440, 1440, 1440] should apply
    to all 3 grid columns and autofit should be disabled.
    """
    doc = declarative.Document(
        sections=[
            declarative.Section(
                children=[
                    declarative.Table(
                        rows=[
                            declarative.TableRow(
                                children=[
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="Wide")],
                                        grid_span=2,
                                    ),
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="Right")],
                                    ),
                                ],
                            ),
                            declarative.TableRow(
                                children=[
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="A")],
                                    ),
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="B")],
                                    ),
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="C")],
                                    ),
                                ],
                            ),
                        ],
                        column_widths=[1440, 1440, 1440],
                    ),
                ],
            ),
        ],
    )

    docx_doc = await doc.to_docx()
    tbl = docx_doc.tables[0]

    assert len(tbl.columns) == 3  # noqa: PLR2004
    assert tbl.autofit is False
    assert tbl.columns[0].width == shared.Twips(1440)
    assert tbl.columns[1].width == shared.Twips(1440)
    assert tbl.columns[2].width == shared.Twips(1440)


@pytest.mark.asyncio
async def test_table_column_widths_mismatch_raises() -> None:
    """Test that a column_widths list of wrong length raises ValueError.

    A table with 1 column but column_widths=[1440, 2880] (2 entries) should
    raise ValueError when the document is rendered.
    """
    doc = declarative.Document(
        sections=[
            declarative.Section(
                children=[
                    declarative.Table(
                        rows=[
                            declarative.TableRow(
                                children=[
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="Only")],
                                    ),
                                ],
                            ),
                        ],
                        column_widths=[1440, 2880],
                    ),
                ],
            ),
        ],
    )

    with pytest.raises(ValueError, match="column_widths length"):
        await doc.to_docx()


@pytest.mark.asyncio
async def test_table_layout_fixed() -> None:
    """Test that layout="fixed" disables autofit without setting column widths.

    A table with layout="fixed" and no column_widths should have autofit=False.
    """
    doc = declarative.Document(
        sections=[
            declarative.Section(
                children=[
                    declarative.Table(
                        rows=[
                            declarative.TableRow(
                                children=[
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="Cell")],
                                    ),
                                ],
                            ),
                        ],
                        layout="fixed",
                    ),
                ],
            ),
        ],
    )

    docx_doc = await doc.to_docx()
    docx_table = docx_doc.tables[0]

    assert docx_table.autofit is False


@pytest.mark.asyncio
async def test_table_layout_autofit_overrides_column_widths() -> None:
    """Test that layout="autofit" overrides column_widths and keeps autofit=True.

    A table with both column_widths and layout="autofit" should have autofit
    enabled and the column widths should NOT be applied.
    """
    doc = declarative.Document(
        sections=[
            declarative.Section(
                children=[
                    declarative.Table(
                        rows=[
                            declarative.TableRow(
                                children=[
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="A")],
                                    ),
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="B")],
                                    ),
                                    declarative.TableCell(
                                        children=[declarative.Paragraph(text="C")],
                                    ),
                                ],
                            ),
                        ],
                        column_widths=[1440, 1440, 1440],
                        layout="autofit",
                    ),
                ],
            ),
        ],
    )

    docx_doc = await doc.to_docx()
    docx_table = docx_doc.tables[0]

    assert docx_table.autofit is True
    assert docx_table.columns[0].width != shared.Twips(1440)
