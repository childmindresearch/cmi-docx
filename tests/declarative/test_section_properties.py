"""Tests for SectionProperties page orientation in the declarative API."""

import pytest
from docx.shared import Inches

from cmi_docx import declarative


@pytest.mark.asyncio
async def test_landscape_no_page_size() -> None:
    """Test landscape orientation without explicit page_size.

    When page_orientation is 'landscape' and no page_size is provided, the
    resulting section should have page_width greater than page_height.
    """
    doc = declarative.Document(
        sections=[
            declarative.Section(
                children=[declarative.Paragraph(text="Landscape page")],
                properties=declarative.SectionProperties(
                    page_orientation="landscape",
                ),
            ),
        ],
    )

    docx_doc = await doc.to_docx()
    section = docx_doc.sections[-1]
    assert section.page_width is not None
    assert section.page_height is not None
    assert section.page_width > section.page_height


@pytest.mark.asyncio
async def test_landscape_with_page_size() -> None:
    """Test landscape orientation with portrait-sized page_size dimensions.

    When page_orientation is 'landscape' and page_size provides portrait
    dimensions (width < height), the dimensions should be swapped so that
    the resulting section has page_width greater than page_height.
    """
    doc = declarative.Document(
        sections=[
            declarative.Section(
                children=[declarative.Paragraph(text="Landscape with size")],
                properties=declarative.SectionProperties(
                    page_orientation="landscape",
                    page_size={"width": Inches(8.5), "height": Inches(11)},
                ),
            ),
        ],
    )

    docx_doc = await doc.to_docx()
    section = docx_doc.sections[-1]
    assert section.page_width is not None
    assert section.page_height is not None
    assert section.page_width > section.page_height


@pytest.mark.asyncio
async def test_portrait_no_page_size() -> None:
    """Test portrait orientation without explicit page_size.

    When page_orientation is 'portrait' and no page_size is provided, the
    resulting section should have page_width less than page_height.
    """
    doc = declarative.Document(
        sections=[
            declarative.Section(
                children=[declarative.Paragraph(text="Portrait page")],
                properties=declarative.SectionProperties(
                    page_orientation="portrait",
                ),
            ),
        ],
    )

    docx_doc = await doc.to_docx()
    section = docx_doc.sections[-1]
    assert section.page_width is not None
    assert section.page_height is not None
    assert section.page_width < section.page_height


@pytest.mark.asyncio
async def test_no_orientation() -> None:
    """Test section without any page_orientation set.

    When no page_orientation is specified, the resulting section should still
    have valid (non-None) page_width and page_height values.
    """
    doc = declarative.Document(
        sections=[
            declarative.Section(
                children=[declarative.Paragraph(text="Default orientation")],
            ),
        ],
    )

    docx_doc = await doc.to_docx()
    section = docx_doc.sections[-1]
    assert section.page_width is not None
    assert section.page_height is not None


@pytest.mark.asyncio
async def test_landscape_multi_section() -> None:
    """Test that landscape properties apply to the correct section index.

    With 2 declarative sections (portrait then landscape), the resulting
    docx_doc.sections list has exactly 2 entries: sections[0] is portrait
    and sections[1] is landscape.
    """
    doc = declarative.Document(
        sections=[
            declarative.Section(
                children=[declarative.Paragraph(text="Portrait page")],
                properties=declarative.SectionProperties(
                    page_orientation="portrait",
                ),
            ),
            declarative.Section(
                children=[declarative.Paragraph(text="Landscape page")],
                properties=declarative.SectionProperties(
                    page_orientation="landscape",
                ),
            ),
        ],
    )

    docx_doc = await doc.to_docx()
    sections = list(docx_doc.sections)

    # 2 declarative sections produce exactly 2 sections.
    expected_section_count = 2
    assert len(sections) == expected_section_count

    portrait_section = sections[0]
    landscape_section = sections[1]

    assert portrait_section.page_width is not None
    assert portrait_section.page_height is not None
    assert portrait_section.page_width < portrait_section.page_height

    assert landscape_section.page_width is not None
    assert landscape_section.page_height is not None
    assert landscape_section.page_width > landscape_section.page_height
