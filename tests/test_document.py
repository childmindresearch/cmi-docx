"""Tests for the document module."""

import docx

from cmi_docx import document


def test_find_in_paragraphs() -> None:
    """Test finding a text in a document's paragraphs."""
    doc = docx.Document()
    doc.add_paragraph("Hello, world!")
    doc.add_paragraph("Hello, world, Hello!")
    extend_document = document.ExtendDocument(doc)

    actual = extend_document.find_in_paragraphs("Hello")

    assert actual[0].character_indices == [(0, 5)]
    assert actual[1].character_indices == [(0, 5), (14, 19)]


def test_find_in_runs() -> None:
    """Test finding a text in a document's runs."""
    doc = docx.Document()
    paragraph = doc.add_paragraph("Hello, world!")
    paragraph.add_run("Hello, world, Hello!")
    extend_document = document.ExtendDocument(doc)

    actual = extend_document.find_in_runs("Hello")

    assert actual[0].run_indices == (0, 0)
    assert actual[0].character_indices == (0, 5)
    assert actual[1].run_indices == (1, 1)
    assert actual[1].character_indices == (0, 5)
    assert actual[2].run_indices == (1, 1)
    assert actual[2].character_indices == (14, 19)


def test_replace() -> None:
    """Test replacing text in a document."""
    doc = docx.Document()
    doc.add_paragraph("Hello, world!")
    extend_document = document.ExtendDocument(doc)

    extend_document.replace("Hello", "Goodbye")

    assert doc.paragraphs[0].text == "Goodbye, world!"


def test_replace_across_runs() -> None:
    """Test replacing text across runs in a document."""
    doc = docx.Document()
    paragraph = doc.add_paragraph("Hello, world!")
    paragraph.add_run(" Maintain, World!")
    paragraph.add_run(" Goodbye, World!")
    extend_document = document.ExtendDocument(doc)

    extend_document.replace("world! Maintain, World! Goodbye", "Goodbye")

    assert doc.paragraphs[0].text == "Hello, Goodbye, World!"


def test_insert_paragraph_by_object() -> None:
    """Test inserting a paragraph into a document."""
    doc = docx.Document()
    doc.add_paragraph("Hello, world!")
    doc.add_paragraph("Goodbye, world!")
    extend_document = document.ExtendDocument(doc)
    new_paragraph = docx.Document().add_paragraph("Maintain, world!")

    extend_document.insert_paragraph_by_object(1, new_paragraph)

    assert doc.paragraphs[0].text == "Hello, world!"
    assert doc.paragraphs[1].text == "Maintain, world!"
    assert doc.paragraphs[2].text == "Goodbye, world!"


def test_insert_paragraph_by_text() -> None:
    """Test inserting a paragraph into a document."""
    doc = docx.Document()
    doc.add_paragraph("Hello, world!")
    doc.add_paragraph("Goodbye, world!")
    extend_document = document.ExtendDocument(doc)

    extend_document.insert_paragraph_by_text(1, "Maintain, world!")

    assert doc.paragraphs[0].text == "Hello, world!"
    assert doc.paragraphs[1].text == "Maintain, world!"
    assert doc.paragraphs[2].text == "Goodbye, world!"
