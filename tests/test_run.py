"""Tests for the run module."""

import docx
import pytest

from cmi_docx import run


def test_find_run_lt_same_paragraph() -> None:
    """Test that comparing FindRun of the same paragraph works."""
    document = docx.Document()
    paragraph = document.add_paragraph("Hello, world!")

    find_run1 = run.FindRun(paragraph, (0, 1), (0, 5))
    find_run2 = run.FindRun(paragraph, (1, 2), (5, 10))

    assert find_run1 < find_run2


def test_find_run_lt_different_paragraphs() -> None:
    """Test that comparing FindRun of different paragraphs fails."""
    document = docx.Document()
    paragraph1 = document.add_paragraph("Hello, world!")
    paragraph2 = document.add_paragraph("Hello, world!")

    find_run1 = run.FindRun(paragraph1, (0, 1), (0, 5))
    find_run2 = run.FindRun(paragraph2, (0, 1), (0, 5))

    with pytest.raises(ValueError):
        assert find_run1 < find_run2


def test_extend_run_format() -> None:
    """Test that formatting a run works."""
    document = docx.Document()
    paragraph = document.add_paragraph("Hello, world!")
    paragraph_run = paragraph.runs[0]

    extend_run = run.ExtendRun(paragraph_run)
    extend_run.format(
        bold=True,
        italics=True,
        underline=True,
        strike=True,
        superscript=True,
        font_rgb=(1, 0, 0),
    )

    assert paragraph_run.bold
    assert paragraph_run.italic
    assert paragraph_run.underline
    assert paragraph_run.strike
    assert paragraph_run.font.superscript
    assert paragraph_run.font.color.rgb == (1, 0, 0)
