"""Tests for the paragraph module."""

import docx
import pytest
from docx.text import paragraph as docx_paragraph

from cmi_docx import paragraph


@pytest.fixture
def sample_paragraph() -> docx_paragraph.Paragraph:
    """Returns a sample paragraph."""
    document = docx.Document()
    return document.add_paragraph("This is a sample paragraph.")


def test_find_single_in_paragraph(sample_paragraph: docx_paragraph.Paragraph) -> None:
    """Test finding a text in a paragraph."""
    extend_paragraph = paragraph.ExtendParagraph(sample_paragraph)
    expected = paragraph.FindParagraph(
        paragraph=sample_paragraph,
        character_indices=[(10, 16)],
    )

    actual = extend_paragraph.find_in_paragraph("sample")

    assert actual == expected


def test_find_multiple_in_paragraph(sample_paragraph: docx_paragraph.Paragraph) -> None:
    """Test finding multiple texts in a paragraph."""
    extend_paragraph = paragraph.ExtendParagraph(sample_paragraph)
    expected = paragraph.FindParagraph(
        paragraph=sample_paragraph,
        character_indices=[(2, 4), (5, 7)],
    )

    actual = extend_paragraph.find_in_paragraph("is")

    assert actual == expected


def test_find_in_single_run(sample_paragraph: docx_paragraph.Paragraph) -> None:
    """Test finding a text in a single paragraph run."""
    extend_paragraph = paragraph.ExtendParagraph(sample_paragraph)
    expected = [
        paragraph.run.FindRun(
            paragraph=sample_paragraph,
            run_indices=(0, 0),
            character_indices=(2, 4),
        ),
        paragraph.run.FindRun(
            paragraph=sample_paragraph,
            run_indices=(0, 0),
            character_indices=(5, 7),
        ),
    ]

    actual = extend_paragraph.find_in_runs("is")

    assert actual[0].paragraph.text == expected[0].paragraph.text
    assert actual[0].run_indices == expected[0].run_indices
    assert actual[0].character_indices == expected[0].character_indices
    assert actual[1].paragraph.text == expected[1].paragraph.text
    assert actual[1].run_indices == expected[1].run_indices
    assert actual[1].character_indices == expected[1].character_indices


def test_replace_single_run(sample_paragraph: docx_paragraph.Paragraph) -> None:
    """Test replacing text in a paragraph."""
    extend_paragraph = paragraph.ExtendParagraph(sample_paragraph)

    extend_paragraph.replace("sample", "example")

    assert sample_paragraph.text == "This is a example paragraph."


def test_replace_multiple_runs(sample_paragraph: docx_paragraph.Paragraph) -> None:
    """Test replacing text in multiple runs."""
    sample_paragraph.add_run(" This is a sample paragraph.")
    extend_paragraph = paragraph.ExtendParagraph(sample_paragraph)

    extend_paragraph.replace("This is", "That was")

    assert (
        sample_paragraph.text
        == "That was a sample paragraph. That was a sample paragraph."
    )


def test_add_styled_runs(sample_paragraph: docx_paragraph.Paragraph) -> None:
    """Test adding styled runs to a paragraph."""
    extend_paragraph = paragraph.ExtendParagraph(sample_paragraph)
    base_runs = len(sample_paragraph.runs)

    text = ["Hello", "World"]
    styles = [{"bold": True}, {"italics": True, "underline": True}]

    extend_paragraph.add_styled_runs(text, styles)

    assert sample_paragraph.runs[base_runs].text == "Hello"
    assert sample_paragraph.runs[base_runs].bold
    assert sample_paragraph.runs[base_runs + 1].text == "World"
    assert sample_paragraph.runs[base_runs + 1].italic
    assert sample_paragraph.runs[base_runs + 1].underline
