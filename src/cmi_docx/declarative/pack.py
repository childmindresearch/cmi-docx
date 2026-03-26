"""Packing logic to convert declarative components to python-docx objects."""

import io
import pathlib
from typing import TYPE_CHECKING

import docx
from docx import document
from docx.enum import text as docx_text
from docx.enum.section import WD_ORIENTATION
from docx.shared import Pt, RGBColor

from cmi_docx import declarative

if TYPE_CHECKING:
    from cmi_docx.declarative.image import ImageRun
    from cmi_docx.declarative.paragraph import Paragraph, TextRun
    from cmi_docx.declarative.section import Footer, Header, Section
    from cmi_docx.declarative.table import Table, TableCell, TableRow


def pack(doc: declarative.Document) -> document.Document:
    """Convert a declarative Document into a python-docx Document.

    Args:
        doc: The declarative Document to convert.

    Returns:
        A python-docx Document ready to be saved.
    """
    docx_doc = docx.Document()

    if doc.creator:
        docx_doc.core_properties.author = doc.creator
    if doc.title:
        docx_doc.core_properties.title = doc.title
    if doc.subject:
        docx_doc.core_properties.subject = doc.subject
    if doc.description:
        docx_doc.core_properties.comments = doc.description
    if doc.keywords:
        docx_doc.core_properties.keywords = doc.keywords
    if doc.category:
        docx_doc.core_properties.category = doc.category

    for section in doc.sections:
        _pack_section(docx_doc, section)

    return docx_doc


def _pack_section(docx_doc: "DocxDocument", section: "Section") -> None:
    """Pack a Section into a python-docx document.

    Args:
        docx_doc: The python-docx Document.
        section: The declarative Section.
    """
    if section.children:
        for child in section.children:
            _pack_block_element(docx_doc, child)

    docx_section = docx_doc.sections[-1]

    if section.properties:
        props = section.properties
        if props.page_size:
            if "width" in props.page_size:
                docx_section.page_width = props.page_size["width"]
            if "height" in props.page_size:
                docx_section.page_height = props.page_size["height"]

        if props.page_margins:
            margin_mapping = {
                "top": "top_margin",
                "bottom": "bottom_margin",
                "left": "left_margin",
                "right": "right_margin",
                "header": "header_distance",
                "footer": "footer_distance",
                "gutter": "gutter",
            }
            for src_key, target_attr in margin_mapping.items():
                if src_key in props.page_margins:
                    setattr(docx_section, target_attr, props.page_margins[src_key])

        if props.page_orientation:
            if props.page_orientation.lower() == "landscape":
                docx_section.orientation = WD_ORIENTATION.LANDSCAPE
            elif props.page_orientation.lower() == "portrait":
                docx_section.orientation = WD_ORIENTATION.PORTRAIT

    if section.headers:
        for header_type, header in section.headers.items():
            _pack_header(docx_section, header_type, header)

    if section.footers:
        for footer_type, footer in section.footers.items():
            _pack_footer(docx_section, footer_type, footer)


def _get_header_or_footer(docx_section: "DocxSection", hf_type: str, is_header: bool):
    """Get the appropriate header or footer from a section.

    Args:
        docx_section: The python-docx Section.
        hf_type: The type ('default', 'first', 'even').
        is_header: True for header, False for footer.

    Returns:
        The header or footer object, or None if type is invalid.
    """
    type_mapping = {
        "default": "header" if is_header else "footer",
        "first": "first_page_header" if is_header else "first_page_footer",
        "even": "even_page_header" if is_header else "even_page_footer",
    }

    attr_name = type_mapping.get(hf_type)
    return getattr(docx_section, attr_name, None) if attr_name else None


def _pack_header(
    docx_section: "DocxSection", header_type: str, header: "Header"
) -> None:
    """Pack a Header into a python-docx section.

    Args:
        docx_section: The python-docx Section.
        header_type: The header type ('default', 'first', 'even').
        header: The declarative Header.
    """
    docx_header = _get_header_or_footer(docx_section, header_type, is_header=True)
    if docx_header and header.children:
        for child in header.children:
            _pack_block_element(docx_header, child)


def _pack_footer(
    docx_section: "DocxSection", footer_type: str, footer: "Footer"
) -> None:
    """Pack a Footer into a python-docx section.

    Args:
        docx_section: The python-docx Section.
        footer_type: The footer type ('default', 'first', 'even').
        footer: The declarative Footer.
    """
    docx_footer = _get_header_or_footer(docx_section, footer_type, is_header=False)
    if docx_footer and footer.children:
        for child in footer.children:
            _pack_block_element(docx_footer, child)


def _pack_block_element(container, element: "Paragraph | Table") -> None:
    """Pack a block-level element (Paragraph or Table).

    Args:
        container: The container to add to (Document, Header, Footer, or Cell).
        element: The Paragraph or Table to pack.
    """
    from cmi_docx.declarative.paragraph import Paragraph
    from cmi_docx.declarative.table import Table

    if isinstance(element, Paragraph):
        _pack_paragraph(container, element)
    elif isinstance(element, Table):
        _pack_table(container, element)


def _pack_paragraph(container, para: "Paragraph") -> None:
    """Pack a Paragraph into a container.

    Args:
        container: The container to add to.
        para: The declarative Paragraph.
    """
    docx_para = container.add_paragraph()

    if para.style:
        docx_para.style = para.style

    if para.heading:
        heading_style_names = {
            1: "Heading 1",
            2: "Heading 2",
            3: "Heading 3",
            4: "Heading 4",
            5: "Heading 5",
            6: "Heading 6",
            7: "Heading 7",
            8: "Heading 8",
            9: "Heading 9",
        }
        if para.heading in heading_style_names:
            docx_para.style = heading_style_names[para.heading]

    if para.alignment:
        docx_para.alignment = para.alignment

    fmt = docx_para.paragraph_format
    if para.spacing_before:
        fmt.space_before = Pt(para.spacing_before)
    if para.spacing_after:
        fmt.space_after = Pt(para.spacing_after)
    if para.line_spacing:
        fmt.line_spacing = para.line_spacing
    if para.left_indent:
        fmt.left_indent = Pt(para.left_indent)
    if para.right_indent:
        fmt.right_indent = Pt(para.right_indent)
    if para.first_line_indent:
        fmt.first_line_indent = Pt(para.first_line_indent)
    if para.keep_together is not None:
        fmt.keep_together = para.keep_together
    if para.keep_with_next is not None:
        fmt.keep_with_next = para.keep_with_next
    if para.page_break_before is not None:
        fmt.page_break_before = para.page_break_before
    if para.widow_control is not None:
        fmt.widow_control = para.widow_control

    if para.text:
        docx_para.add_run(para.text)
    elif para.children:
        for child in para.children:
            _pack_inline_element(docx_para, child)


def _pack_inline_element(docx_para: "DocxParagraph", element) -> None:
    """Pack an inline element (TextRun, ImageRun, Tab, Break).

    Args:
        docx_para: The python-docx Paragraph.
        element: The inline element to pack.
    """
    from cmi_docx.declarative.image import ImageRun
    from cmi_docx.declarative.paragraph import Break, Tab, TextRun

    if isinstance(element, TextRun):
        _pack_text_run(docx_para, element)
    elif isinstance(element, ImageRun):
        _pack_image_run(docx_para, element)
    elif isinstance(element, Tab):
        docx_para.add_run().add_tab()
    elif isinstance(element, Break):
        break_type_mapping = {
            "page": docx_text.WD_BREAK.PAGE,
            "column": docx_text.WD_BREAK.COLUMN,
        }
        break_type = break_type_mapping.get(element.type)
        if break_type:
            docx_para.add_run().add_break(break_type)
        else:
            docx_para.add_run().add_break()


def _pack_text_run(docx_para: "DocxParagraph", run: "TextRun") -> None:
    """Pack a TextRun into a paragraph.

    Args:
        docx_para: The python-docx Paragraph.
        run: The declarative TextRun.
    """
    docx_run = docx_para.add_run(run.text)

    if run.bold is not None:
        docx_run.bold = run.bold
    if run.italic is not None:
        docx_run.italic = run.italic
    if run.underline is not None:
        docx_run.underline = run.underline

    font = docx_run.font
    if run.font:
        font.name = run.font
    if run.size:
        font.size = Pt(run.size)
    if run.color:
        font.color.rgb = RGBColor(*run.color)
    if run.superscript:
        font.superscript = True
    if run.subscript:
        font.subscript = True
    if run.strike:
        font.strike = True
    if run.all_caps:
        font.all_caps = True
    if run.small_caps:
        font.small_caps = True
    if run.highlight:
        font.highlight_color = run.highlight


def _pack_image_run(docx_para: "DocxParagraph", image: "ImageRun") -> None:
    """Pack an ImageRun into a paragraph.

    Args:
        docx_para: The python-docx Paragraph.
        image: The declarative ImageRun.
    """
    docx_run = docx_para.add_run()

    width = None
    height = None
    if image.transformation:
        if "width" in image.transformation:
            width = Pt(image.transformation["width"])
        if "height" in image.transformation:
            height = Pt(image.transformation["height"])

    if isinstance(image.data, bytes):
        docx_run.add_picture(io.BytesIO(image.data), width=width, height=height)
    elif isinstance(image.data, (str, pathlib.Path)):
        docx_run.add_picture(str(image.data), width=width, height=height)


def _pack_table(container, table: "Table") -> None:
    """Pack a Table into a container.

    Args:
        container: The container to add to.
        table: The declarative Table.
    """
    num_rows = len(table.rows)
    num_cols = max(len(row.children) for row in table.rows) if table.rows else 0

    docx_table = container.add_table(rows=num_rows, cols=num_cols)

    if table.style:
        docx_table.style = table.style

    for row_idx, row in enumerate(table.rows):
        _pack_table_row(docx_table.rows[row_idx], row)


def _pack_table_row(docx_row: "DocxRow", row: "TableRow") -> None:
    """Pack a TableRow.

    Args:
        docx_row: The python-docx table row.
        row: The declarative TableRow.
    """
    if row.height and "value" in row.height:
        docx_row.height = row.height["value"]

    if row.cant_split is not None:
        docx_row.cant_split = row.cant_split

    for cell_idx, cell in enumerate(row.children):
        _pack_table_cell(docx_row.cells[cell_idx], cell)


def _pack_table_cell(docx_cell: "DocxCell", cell: "TableCell") -> None:
    """Pack a TableCell.

    Args:
        docx_cell: The python-docx table cell.
        cell: The declarative TableCell.
    """
    if cell.width and "size" in cell.width:
        docx_cell.width = cell.width["size"]

    if cell.children:
        docx_cell.text = ""
        for child in cell.children:
            _pack_block_element(docx_cell, child)
