"""Top-level Document class for declarative API."""

import asyncio
import dataclasses
import datetime
import io
import pathlib
from collections.abc import Sequence

import docx
from docx import document as docx_document
from docx import oxml, shared
from docx import section as docx_section
from docx import table as docx_table
from docx.enum import section as docx_enum_section
from docx.enum import style as docx_style
from docx.enum import text as docx_text
from docx.oxml import simpletypes as docx_simpletypes
from docx.oxml.ns import qn
from docx.text import paragraph as docx_paragraph
from lxml import (
    etree,  # ty:ignore[unresolved-import] # This does work; not sure why not detected.
)

from cmi_docx import document as imperative_document
from cmi_docx.declarative import image, paragraph, section, table
from cmi_docx.declarative import styles as styles_mod


@dataclasses.dataclass
class DocumentTemplate:
    """Defines a template to use as a base for a document.

    Attributes:
        path: Path to the document file.
        replacements: Dictionary of needle/replacements to make in the template.
        paragraph_index: Paragraph index in the template at which to insert
            user-defined section content. If None, content is appended to the end.
            The index refers to the paragraph position in the original template
            (before any insertions).
    """

    path: pathlib.Path | str
    replacements: dict[str, str] | None = None
    paragraph_index: int | None = None


class Document:
    """A Word document with sections.

    This is the top-level container for a declarative document. All operations
    are async - use await to resolve all async children concurrently.

    Attributes:
        sections: List of Section components or coroutines that resolve to sections.
        creator: Document creator metadata.
        title: Document title metadata.
        subject: Document subject metadata.
        description: Document description metadata.
        keywords: Document keywords metadata.
        category: Document category metadata.
        comments: Document comments metadata.
        comment_author: Default author name for all Word comments. Can be overridden
            per-paragraph or per-text-run.
        styles: Document-level style definitions to create or modify.
        numbering: Document-level numbering definitions.

    Example:
        >>> async def create_doc():
        ...     doc = Document(sections=[
        ...         Section(children=[
        ...             Paragraph(text="Hello World"),
        ...             fetch_paragraph(),  # async function
        ...         ]),
        ...     ])
        ...     await doc.save("output.docx")
    """

    def __init__(  # noqa: PLR0913, D107
        self,
        sections: list[section.Section],
        creator: str | None = None,
        title: str | None = None,
        subject: str | None = None,
        description: str | None = None,
        keywords: str | None = None,
        category: str | None = None,
        comments: str | None = None,
        version: str | None = None,
        comment_author: str | None = None,
        styles: (
            list[styles_mod.ParagraphStyleDefinition | styles_mod.TableStyleDefinition]
            | None
        ) = None,
        numbering: dict[str, str | int | list[dict[str, str | int]]] | None = None,
    ) -> None:
        self.sections = sections
        self.creator = creator
        self.title = title
        self.subject = subject
        self.description = description
        self.keywords = keywords
        self.category = category
        self.version = version
        self.comments = comments
        self.comment_author = comment_author
        self.styles = styles
        self.numbering = numbering

    async def to_docx(  # noqa: C901
        self, template: DocumentTemplate | None = None
    ) -> docx_document.Document:
        """Convert to a python-docx Document.

        Automatically resolves all async children before converting.

        Returns:
            A python-docx Document object.
        """
        await asyncio.gather(*(section.resolve() for section in self.sections))

        docx_doc = (
            docx.Document() if template is None else docx.Document(str(template.path))
        )

        if template is not None and template.replacements is not None:
            extended_doc = imperative_document.ExtendDocument(docx_doc)
            for needle, replacement in template.replacements.items():
                extended_doc.replace(needle, replacement)

        if self.creator:
            docx_doc.core_properties.author = self.creator
        if self.title:
            docx_doc.core_properties.title = self.title
        if self.subject:
            docx_doc.core_properties.subject = self.subject
        if self.description:
            docx_doc.core_properties.comments = self.description
        if self.keywords:
            docx_doc.core_properties.keywords = self.keywords
        if self.category:
            docx_doc.core_properties.category = self.category
        if self.version:
            docx_doc.core_properties.version = self.version
        if self.comments:
            docx_doc.core_properties.comments = self.comments
        now = datetime.datetime.now(datetime.UTC)
        docx_doc.core_properties.created = now
        docx_doc.core_properties.modified = now

        if self.styles:
            _apply_style_definitions(docx_doc, self.styles)

        paragraph_index = template.paragraph_index if template is not None else None
        insertion_offset = 0
        for i, sec in enumerate(self.sections):
            current_index = (
                (paragraph_index + insertion_offset)
                if paragraph_index is not None
                else None
            )
            elements_inserted = _pack_section(
                docx_doc,
                sec,
                self.comment_author,
                current_index,
                is_last=(i == len(self.sections) - 1),
            )
            insertion_offset += elements_inserted

        return docx_doc


def _apply_style_definitions(
    docx_doc: docx_document.Document,
    style_definitions: list[
        styles_mod.ParagraphStyleDefinition | styles_mod.TableStyleDefinition
    ],
) -> None:
    """Apply a list of style definitions to a python-docx document.

    Iterates over each definition and either creates or modifies the
    corresponding style in the document.

    Args:
        docx_doc: The python-docx Document to apply styles to.
        style_definitions: List of paragraph or table style definitions.
    """
    for defn in style_definitions:
        if isinstance(defn, styles_mod.ParagraphStyleDefinition):
            _apply_paragraph_style_definition(docx_doc, defn)
        elif isinstance(defn, styles_mod.TableStyleDefinition):
            _apply_table_style_definition(docx_doc, defn)


def _apply_paragraph_style_definition(  # noqa: C901, PLR0912, PLR0915
    docx_doc: docx_document.Document,
    defn: styles_mod.ParagraphStyleDefinition,
) -> None:
    """Apply a ParagraphStyleDefinition to a python-docx document.

    Uses a get-or-create pattern: if the style already exists it is modified
    in place; otherwise a new style is created (with base_style applied).

    Args:
        docx_doc: The python-docx Document to apply styles to.
        defn: The paragraph style definition.
    """
    try:
        style = docx_doc.styles[defn.name]
    except KeyError:
        style = docx_doc.styles.add_style(defn.name, docx_style.WD_STYLE_TYPE.PARAGRAPH)
        if defn.base_style:
            style.base_style = docx_doc.styles[defn.base_style]

    if defn.next_paragraph_style is not None:
        style.next_paragraph_style = docx_doc.styles[defn.next_paragraph_style]

    font = style.font
    if defn.font is not None:
        font.name = defn.font
        # font.element is the <w:style> element; <w:rFonts> lives inside <w:rPr>.
        # We must search the nested <w:rPr> to locate it.
        rpr = font.element.find(qn("w:rPr"))
        r_fonts = rpr.find(qn("w:rFonts")) if rpr is not None else None
        if r_fonts is None:
            # Built-in styles may already have <w:rPr> but no <w:rFonts>;
            # create and insert it as the first child of <w:rPr>.
            if rpr is None:
                rpr = etree.SubElement(font.element, qn("w:rPr"))
            r_fonts = etree.SubElement(rpr, qn("w:rFonts"))
            rpr.insert(0, r_fonts)
        r_fonts.set(qn("w:ascii"), defn.font)
        r_fonts.set(qn("w:hAnsi"), defn.font)
        r_fonts.attrib.pop(qn("w:asciiTheme"), None)
        r_fonts.attrib.pop(qn("w:hAnsiTheme"), None)
    if defn.font_size is not None:
        font.size = shared.Pt(defn.font_size)
    if defn.bold is not None:
        font.bold = defn.bold
    if defn.italic is not None:
        font.italic = defn.italic
    if defn.underline is not None:
        font.underline = defn.underline
    if defn.color is not None:
        font.color.rgb = shared.RGBColor(*defn.color)

    pf = style.paragraph_format
    if defn.alignment is not None:
        pf.alignment = defn.alignment
    if defn.spacing_before is not None:
        pf.space_before = shared.Pt(defn.spacing_before)
    if defn.spacing_after is not None:
        pf.space_after = shared.Pt(defn.spacing_after)
    if defn.line_spacing is not None:
        pf.line_spacing = defn.line_spacing
    if defn.left_indent is not None:
        pf.left_indent = shared.Inches(defn.left_indent)
    if defn.right_indent is not None:
        pf.right_indent = shared.Inches(defn.right_indent)
    if defn.first_line_indent is not None:
        pf.first_line_indent = shared.Inches(defn.first_line_indent)
    if defn.keep_together is not None:
        pf.keep_together = defn.keep_together
    if defn.keep_with_next is not None:
        pf.keep_with_next = defn.keep_with_next
    if defn.page_break_before is not None:
        pf.page_break_before = defn.page_break_before
    if defn.widow_control is not None:
        pf.widow_control = defn.widow_control


def _apply_table_style_definition(  # noqa: C901, PLR0912
    docx_doc: docx_document.Document,
    defn: styles_mod.TableStyleDefinition,
) -> None:
    """Apply a TableStyleDefinition to a python-docx document.

    Always creates a new table style. Per-section formatting is applied via
    ``<w:tblStylePr>`` XML elements appended to the style element.

    Args:
        docx_doc: The python-docx Document to apply styles to.
        defn: The table style definition.
    """
    style = docx_doc.styles.add_style(defn.name, docx_style.WD_STYLE_TYPE.TABLE)
    if defn.base_style:
        style.base_style = docx_doc.styles[defn.base_style]

    if defn.whole_table is not None:
        fmt = defn.whole_table
        font = style.font
        if fmt.font is not None:
            font.name = fmt.font
        if fmt.font_size is not None:
            font.size = shared.Pt(fmt.font_size)
        if fmt.bold is not None:
            font.bold = fmt.bold
        if fmt.italic is not None:
            font.italic = fmt.italic
        if fmt.underline is not None:
            font.underline = fmt.underline
        if fmt.color is not None:
            font.color.rgb = shared.RGBColor(*fmt.color)

        pf = style.paragraph_format
        if fmt.alignment is not None:
            pf.alignment = fmt.alignment
        if fmt.spacing_before is not None:
            pf.space_before = shared.Pt(fmt.spacing_before)
        if fmt.spacing_after is not None:
            pf.space_after = shared.Pt(fmt.spacing_after)

    _SECTION_TYPE_MAP = {  # noqa: N806
        "first_row": "firstRow",
        "last_row": "lastRow",
        "first_column": "firstCol",
        "last_column": "lastCol",
        "banding_1_row": "band1Horz",
        "banding_2_row": "band2Horz",
        "banding_1_column": "band1Vert",
        "banding_2_column": "band2Vert",
        "top_left_cell": "nwCell",
        "top_right_cell": "neCell",
        "bottom_left_cell": "swCell",
        "bottom_right_cell": "seCell",
    }
    for field_name, section_type in _SECTION_TYPE_MAP.items():
        fmt = getattr(defn, field_name)
        if fmt is not None:
            style.element.append(_build_tbl_style_pr(section_type, fmt))


def _build_tbl_style_pr(  # noqa: C901, PLR0912, PLR0915
    section_type: str,
    fmt: styles_mod.TableSectionFormat,
) -> etree._Element:  # type: ignore[name-defined]
    """Build a ``<w:tblStylePr>`` XML element for a table style section.

    Args:
        section_type: The OOXML ``w:type`` value (e.g. ``"firstRow"``).
        fmt: The formatting to apply to this section.

    Returns:
        A ``<w:tblStylePr>`` lxml element ready to be appended to a style.
    """
    tbl_style_pr = oxml.OxmlElement("w:tblStylePr")
    tbl_style_pr.set(qn("w:type"), section_type)

    has_rpr = any(
        [
            fmt.font is not None,
            fmt.font_size is not None,
            fmt.bold is not None,
            fmt.italic is not None,
            fmt.underline is not None,
            fmt.color,
        ]
    )
    if has_rpr:
        r_pr = oxml.OxmlElement("w:rPr")
        if fmt.bold is not None:
            b = oxml.OxmlElement("w:b")
            if not fmt.bold:
                b.set(qn("w:val"), "0")
            r_pr.append(b)
        if fmt.italic is not None:
            i = oxml.OxmlElement("w:i")
            if not fmt.italic:
                i.set(qn("w:val"), "0")
            r_pr.append(i)
        if fmt.underline is not None:
            u = oxml.OxmlElement("w:u")
            u.set(qn("w:val"), "single" if fmt.underline else "none")
            r_pr.append(u)
        if fmt.color:
            color_el = oxml.OxmlElement("w:color")
            color_el.set(
                qn("w:val"),
                f"{fmt.color[0]:02X}{fmt.color[1]:02X}{fmt.color[2]:02X}",
            )
            r_pr.append(color_el)
        if fmt.font is not None:
            r_fonts = oxml.OxmlElement("w:rFonts")
            r_fonts.set(qn("w:ascii"), fmt.font)
            r_fonts.set(qn("w:hAnsi"), fmt.font)
            r_pr.append(r_fonts)
        if fmt.font_size is not None:
            sz = oxml.OxmlElement("w:sz")
            sz.set(qn("w:val"), str(fmt.font_size * 2))  # half-points
            r_pr.append(sz)
        tbl_style_pr.append(r_pr)

    has_ppr = any(
        [
            fmt.alignment is not None,
            fmt.spacing_before is not None,
            fmt.spacing_after is not None,
        ]
    )
    if has_ppr:
        p_pr = oxml.OxmlElement("w:pPr")
        if fmt.alignment is not None:
            jc = oxml.OxmlElement("w:jc")
            alignment_map = {
                0: "left",
                1: "center",
                2: "right",
                3: "both",
                4: "distribute",
                5: "mediumKashida",
                7: "highKashida",
                8: "lowKashida",
                9: "thaiDistribute",
            }
            jc.set(qn("w:val"), alignment_map.get(int(fmt.alignment), "left"))
            p_pr.append(jc)
        if fmt.spacing_before is not None or fmt.spacing_after is not None:
            spacing = oxml.OxmlElement("w:spacing")
            if fmt.spacing_before is not None:
                spacing.set(qn("w:before"), str(int(fmt.spacing_before * 20)))
            if fmt.spacing_after is not None:
                spacing.set(qn("w:after"), str(int(fmt.spacing_after * 20)))
            p_pr.append(spacing)
        tbl_style_pr.append(p_pr)

    if fmt.background:
        tc_pr = oxml.OxmlElement("w:tcPr")
        shd = oxml.OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(
            qn("w:fill"),
            f"{fmt.background[0]:02X}{fmt.background[1]:02X}{fmt.background[2]:02X}",
        )
        tc_pr.append(shd)
        tbl_style_pr.append(tc_pr)

    return tbl_style_pr


def _pack_section(  # noqa: C901, PLR0912
    docx_doc: docx_document.Document,
    sec: section.Section,
    default_comment_author: str | None,
    paragraph_index: int | None = None,
    *,
    is_last: bool = False,
) -> int:
    """Pack a Section into a python-docx document.

    Args:
        docx_doc: The python-docx Document.
        sec: The declarative Section.
        default_comment_author: Default author for comments.
        paragraph_index: If provided, insert children starting at this paragraph
            index instead of appending.
        is_last: If True, skip adding a new section at the end (avoids a
            trailing blank page after the final section).

    Returns:
        The number of block elements inserted (for offset tracking).
    """
    if not sec.condition():
        return 0

    paragraphs_inserted = 0
    if sec.children:
        for child in sec.children:  # ty:ignore[not-iterable] callables have been resolved.
            current_index = (
                (paragraph_index + paragraphs_inserted)
                if paragraph_index is not None
                else None
            )
            _pack_block_element(
                docx_doc,
                docx_doc,
                child,  # ty:ignore[invalid-argument-type] already awaited.
                default_comment_author,
                current_index,
            )
            if child.condition() and isinstance(child, paragraph.Paragraph):  # ty:ignore[unresolved-attribute] already awaited.
                paragraphs_inserted += 1

    current_section: docx_section.Section = docx_doc.sections[-1]

    if sec.properties:
        props = sec.properties
        if props.page_size or props.page_orientation:
            if props.page_size:
                width = props.page_size.get("width") or current_section.page_width
                height = props.page_size.get("height") or current_section.page_height
            else:
                width = current_section.page_width
                height = current_section.page_height

            if props.page_orientation:
                orientation = props.page_orientation.lower()
                if orientation == "landscape":
                    if width < height:  # ty:ignore[unsupported-operator]
                        width, height = height, width
                    current_section.page_width = width  # ty:ignore[invalid-assignment]
                    current_section.page_height = height  # ty:ignore[invalid-assignment]
                    current_section.orientation = (
                        docx_enum_section.WD_ORIENTATION.LANDSCAPE
                    )
                elif orientation == "portrait":
                    if width > height:  # ty:ignore[unsupported-operator]
                        width, height = height, width
                    current_section.page_width = width  # ty:ignore[invalid-assignment]
                    current_section.page_height = height  # ty:ignore[invalid-assignment]
                    current_section.orientation = (
                        docx_enum_section.WD_ORIENTATION.PORTRAIT
                    )
            else:
                current_section.page_width = width  # ty:ignore[invalid-assignment]
                current_section.page_height = height  # ty:ignore[invalid-assignment]

        if props.page_margins:
            margin_attrs = {
                "top": "top_margin",
                "bottom": "bottom_margin",
                "left": "left_margin",
                "right": "right_margin",
            }
            for key, attr in margin_attrs.items():
                value = props.page_margins.get(key)  # ty:ignore[invalid-argument-type] key is always a valid Literal.
                if value is not None:
                    setattr(current_section, attr, value)

    if sec.headers:
        for header_type, header in sec.headers.items():
            _pack_header(
                docx_doc, current_section, header_type, header, default_comment_author
            )

    if sec.footers:
        for footer_type, footer in sec.footers.items():
            _pack_footer(
                docx_doc, current_section, footer_type, footer, default_comment_author
            )

    if not is_last:
        docx_doc.add_section()

    return paragraphs_inserted


def _get_header_or_footer(
    section: docx_section.Section, hf_type: str, *, is_header: bool
) -> docx_section.Section | None:
    """Get the appropriate header or footer from a section.

    Args:
        section: The python-docx Section.
        hf_type: The type ('default', 'first', 'even').
        is_header: True for header, False for footer.

    Returns:
        The header or footer object, or None if type is invalid.
    """
    type_mapping = {
        "default": "header" if is_header else "footer",
        "first": "first_page_header" if is_header else "first_page_footer",
        "even": "even_page_header" if is_header else "even_page_footer",
    }

    attr_name = type_mapping.get(hf_type)
    return getattr(section, attr_name, None) if attr_name else None


def _pack_header(
    docx_doc: docx_document.Document,
    section: docx_section.Section,
    header_type: str,
    header: section.Header,
    default_comment_author: str | None,
) -> None:
    """Pack a Header into a python-docx section.

    Args:
        docx_doc: The python-docx Document.
        section: The python-docx Section.
        header_type: The header type ('default', 'first', 'even').
        header: The declarative Header.
        default_comment_author: Default author for comments.
    """
    if not header.condition():
        return

    docx_header = _get_header_or_footer(section, header_type, is_header=True)
    if docx_header and header.children:
        for child in header.children:  # ty:ignore[not-iterable] callables have been resolved.
            _pack_block_element(docx_doc, docx_header, child, default_comment_author)  # ty:ignore[invalid-argument-type] already awaited.


def _pack_footer(
    docx_doc: docx_document.Document,
    section: docx_section.Section,
    footer_type: str,
    footer: section.Footer,
    default_comment_author: str | None,
) -> None:
    """Pack a Footer into a python-docx section.

    Args:
        docx_doc: The python-docx Document.
        section: The python-docx Section.
        footer_type: The footer type ('default', 'first', 'even').
        footer: The declarative Footer.
        default_comment_author: Default author for comments.
    """
    if not footer.condition():
        return

    docx_footer = _get_header_or_footer(section, footer_type, is_header=False)
    if docx_footer and footer.children:
        for child in footer.children:  # ty:ignore[not-iterable] callables have been resolved.
            _pack_block_element(docx_doc, docx_footer, child, default_comment_author)  # ty:ignore[invalid-argument-type] already awaited.


def _pack_block_element(
    docx_doc: docx_document.Document,
    container: docx_document.Document,
    element: paragraph.Paragraph | table.Table,
    default_comment_author: str | None,
    insert_index: int | None = None,
) -> None:
    """Pack a block-level element (Paragraph or Table).

    Args:
        docx_doc: The python-docx Document (needed for comment API).
        container: The container to add to.
        element: The Paragraph or Table to pack.
        default_comment_author: Default author for comments.
        insert_index: If provided, insert at this paragraph index in docx_doc
            instead of appending.
    """
    if not element.condition():
        return None

    if isinstance(element, paragraph.Paragraph):
        return _pack_paragraph(
            docx_doc, container, element, default_comment_author, insert_index
        )
    return _pack_table(
        docx_doc, container, element, default_comment_author, insert_index
    )


def _pack_paragraph(
    docx_doc: docx_document.Document,
    container: docx_document.Document,
    para: paragraph.Paragraph,
    default_comment_author: str | None,
    insert_index: int | None = None,
) -> None:
    """Pack a Paragraph into a container.

    Args:
        docx_doc: The python-docx Document (needed for comment API).
        container: The container to add to.
        para: The declarative Paragraph.
        default_comment_author: Default author for comments.
        insert_index: If provided, insert at this paragraph index in docx_doc
            instead of appending.
    """
    if insert_index is not None:
        n_paragraphs = len(docx_doc.paragraphs)
        if insert_index >= n_paragraphs:
            docx_para = container.add_paragraph()
        else:
            docx_para = docx_doc.paragraphs[insert_index]._insert_paragraph_before()  # noqa: SLF001
    else:
        docx_para = container.add_paragraph()
    _pack_paragraph_into_existing(docx_doc, docx_para, para, default_comment_author)


def _pack_paragraph_into_existing(  # noqa: C901, PLR0912
    docx_doc: docx_document.Document,
    docx_para: docx_paragraph.Paragraph,
    para: paragraph.Paragraph,
    default_comment_author: str | None,
) -> None:
    """Pack a Paragraph into an existing python-docx paragraph.

    Args:
        docx_doc: The python-docx Document (needed for comment API).
        docx_para: The python-docx Paragraph to populate.
        para: The declarative Paragraph.
        default_comment_author: Default author for comments.
    """
    if para.style:
        docx_para.style = para.style

    if para.heading:
        docx_para.style = f"Heading {para.heading}"

    if para.alignment is not None:
        docx_para.alignment = para.alignment

    fmt = docx_para.paragraph_format
    if para.spacing_before is not None:
        fmt.space_before = shared.Pt(para.spacing_before)
    if para.spacing_after is not None:
        fmt.space_after = shared.Pt(para.spacing_after)
    if para.line_spacing is not None:
        fmt.line_spacing = para.line_spacing
    if para.left_indent is not None:
        fmt.left_indent = shared.Pt(para.left_indent)
    if para.right_indent is not None:
        fmt.right_indent = shared.Pt(para.right_indent)
    if para.first_line_indent is not None:
        fmt.first_line_indent = shared.Pt(para.first_line_indent)
    if para.keep_together is not None:
        fmt.keep_together = para.keep_together
    if para.keep_with_next is not None:
        fmt.keep_with_next = para.keep_with_next
    if para.page_break_before is not None:
        fmt.page_break_before = para.page_break_before
    if para.widow_control is not None:
        fmt.widow_control = para.widow_control

    if para.text:
        docx_para.add_run(para.text)  # ty:ignore[invalid-argument-type] Text is already awaited.
    elif para.children:
        for child in para.children:  # ty:ignore[not-iterable] callables have been resolved.
            _pack_inline_element(docx_doc, docx_para, child, default_comment_author)  # ty:ignore[invalid-argument-type] already awaited.

    if para.comment_text:
        author = para.comment_author or default_comment_author or ""
        docx_doc.add_comment(runs=docx_para.runs, text=para.comment_text, author=author)  # ty:ignore[invalid-argument-type] already awaited.


def _pack_inline_element(
    docx_doc: docx_document.Document,
    para: docx_paragraph.Paragraph,
    element: paragraph.TextRun | image.ImageRun | paragraph.Tab | paragraph.Break,
    default_comment_author: str | None,
) -> None:
    """Pack an inline element (TextRun, ImageRun, Tab, Break).

    Args:
        docx_doc: The python-docx Document (needed for comment API).
        para: The python-docx Paragraph.
        element: The inline element to pack.
        default_comment_author: Default author for comments.
    """
    if not element.condition():
        return

    if isinstance(element, paragraph.TextRun):
        _pack_text_run(docx_doc, para, element, default_comment_author)
    elif isinstance(element, image.ImageRun):
        _pack_image_run(para, element)
    elif isinstance(element, paragraph.Tab):
        para.add_run().add_tab()
    elif isinstance(element, paragraph.Break):
        break_type_mapping = {
            "page": docx_text.WD_BREAK.PAGE,
            "column": docx_text.WD_BREAK.COLUMN,
        }
        break_type = break_type_mapping.get(element.type)
        if break_type:
            para.add_run().add_break(break_type)
        else:
            para.add_run().add_break()


def _pack_text_run(  # noqa: C901
    docx_doc: docx_document.Document,
    para: docx_paragraph.Paragraph,
    run: paragraph.TextRun,
    default_comment_author: str | None,
) -> None:
    """Pack a TextRun into a paragraph.

    Args:
        docx_doc: The python-docx Document (needed for comment API).
        para: The python-docx Paragraph.
        run: The declarative TextRun.
        default_comment_author: Default author for comments.
    """
    docx_run = para.add_run(run.text)  # ty:ignore[invalid-argument-type] Already awaited.

    if run.bold is not None:
        docx_run.bold = run.bold
    if run.italic is not None:
        docx_run.italic = run.italic
    if run.underline is not None:
        docx_run.underline = run.underline

    font = docx_run.font
    if run.font:
        font.name = run.font
    if run.size:
        font.size = shared.Pt(run.size)
    if run.color:
        font.color.rgb = shared.RGBColor(*run.color)
    if run.superscript:
        font.superscript = True
    if run.subscript:
        font.subscript = True
    if run.strike:
        font.strike = True
    if run.all_caps:
        font.all_caps = True
    if run.small_caps:
        font.small_caps = True

    if run.comment_text:
        author = run.comment_author or default_comment_author or ""
        docx_doc.add_comment(runs=docx_run, text=run.comment_text, author=author)  # ty:ignore[invalid-argument-type] already awaited.


def _pack_image_run(para: docx_paragraph.Paragraph, img: image.ImageRun) -> None:
    """Pack an ImageRun into a paragraph.

    Args:
        para: The python-docx Paragraph.
        img: The declarative ImageRun.
    """
    docx_run = para.add_run()

    width = None
    height = None
    if img.transformation:
        if "width" in img.transformation:
            width = shared.Pt(img.transformation["width"])
        if "height" in img.transformation:
            height = shared.Pt(img.transformation["height"])

    if isinstance(img.data, bytes):
        docx_run.add_picture(
            io.BytesIO(img.data),
            width=width,
            height=height,
        )
    elif isinstance(img.data, (str, pathlib.Path)):
        docx_run.add_picture(
            str(img.data),
            width=width,
            height=height,
        )


def _pack_table(
    docx_doc: docx_document.Document,
    container: docx_document.Document,
    tbl: table.Table,
    default_comment_author: str | None,
    insert_index: int | None = None,
) -> None:
    """Pack a Table into a container.

    Args:
        docx_doc: The python-docx Document (needed for comment API).
        container: The container to add to.
        tbl: The declarative Table.
        default_comment_author: Default author for comments.
        insert_index: If provided, insert the table before this paragraph index
            in docx_doc instead of appending.
    """
    filtered_rows = [row for row in tbl.rows if row.condition()]  # ty:ignore[unresolved-attribute] already awaited. # ty:ignore[not-iterable] callables have been resolved.

    if not filtered_rows:
        return

    num_rows = len(filtered_rows)
    num_cols = (
        max(
            sum(
                (cell.grid_span or 1)  # ty:ignore[unresolved-attribute] already awaited.
                for cell in row.children  # ty:ignore[not-iterable, unresolved-attribute] callables have been resolved.
                if cell.condition()  # ty:ignore[unresolved-attribute] already awaited.
            )
            for row in filtered_rows
        )
        if filtered_rows
        else 0
    )

    docx_table = container.add_table(rows=num_rows, cols=num_cols)

    if insert_index is not None:
        n_paragraphs = len(docx_doc.paragraphs)
        if insert_index < n_paragraphs:
            target_element = docx_doc.paragraphs[insert_index]._element  # noqa: SLF001
            target_element.addprevious(docx_table._tbl)  # noqa: SLF001

    if tbl.style:
        docx_table.style = tbl.style

    if tbl.layout == "autofit":
        docx_table.autofit = True
    elif tbl.layout == "fixed" or tbl.column_widths is not None:
        # column_widths or layout="fixed" both imply fixed layout
        docx_table.autofit = False

    for row_idx, row in enumerate(filtered_rows):
        _pack_table_row(docx_doc, docx_table.rows[row_idx], row, default_comment_author)  # ty:ignore[invalid-argument-type] already awaited.

    if tbl.layout != "autofit" and tbl.column_widths is not None:
        _apply_column_widths(docx_table, tbl.column_widths, num_cols)


def _apply_column_widths(
    tbl: docx_table.Table,
    column_widths: Sequence[int],
    num_cols: int,
) -> None:
    """Apply fixed column widths to a python-docx table.

    Word XML Elements & Attributes Referenced:
        w:tblW: Table Width property. Defines the overall width of the entire table.
        w:w: Width value attribute. Holds the numeric measurement for an element.
        w:type: Measurement unit type.
        w:tr: Table Row element.
        w:tc: Table Cell element.
        w:gridSpan: Grid span property. Dictates how many logical columns a horizontally
                    merged cell spans.
        w:val: Value attribute.
        w:tcW: Table Cell Width property.

    Args:
        tbl: The python-docx Table to modify.
        column_widths: List of widths in twips (DXA), one per column.
        num_cols: Expected number of columns (must match len(column_widths)).

    Raises:
        ValueError: If len(column_widths) != num_cols.
    """
    if len(column_widths) != num_cols:
        msg = (
            f"column_widths length ({len(column_widths)}) "
            f"must match number of columns ({num_cols})"
        )
        raise ValueError(msg)

    for index, width in enumerate(column_widths):
        tbl.columns[index].width = shared.Twips(width)

    total_width = sum(column_widths)
    tblW = tbl._tbl.tblPr.find(qn("w:tblW"))  # noqa: SLF001, N806
    if tblW is not None:
        tblW.set(qn("w:w"), str(total_width))
        tblW.set(qn("w:type"), "dxa")

    for tr in tbl._tbl.findall(qn("w:tr")):  # noqa: SLF001
        col_start = 0
        for tc in tr.findall(qn("w:tc")):
            gridSpan_el = (  # noqa: N806
                tc.tcPr.find(qn("w:gridSpan")) if tc.tcPr is not None else None
            )
            span = int(gridSpan_el.get(qn("w:val"))) if gridSpan_el is not None else 1
            cell_width = sum(column_widths[col_start : col_start + span])
            tcW = tc.tcPr.find(qn("w:tcW")) if tc.tcPr is not None else None  # noqa: N806
            if tcW is not None:
                tcW.set(qn("w:w"), str(cell_width))
                tcW.set(qn("w:type"), "dxa")
            col_start += span


def _pack_table_row(
    docx_doc: docx_document.Document,
    docx_row: docx_table._Row,
    row: table.TableRow,
    default_comment_author: str | None,
) -> None:
    """Pack a TableRow.

    Args:
        docx_doc: The python-docx Document.
        docx_row: The python-docx table row.
        row: The declarative TableRow.
        default_comment_author: Default author for comments.
    """
    filtered_cells = [cell for cell in row.children if cell.condition()]  # ty:ignore[unresolved-attribute] already awaited. # ty:ignore[not-iterable] callables have been resolved.
    physical_tcs = docx_row._tr.findall(qn("w:tc"))  # noqa: SLF001
    col_offset = 0
    for cell in filtered_cells:
        docx_cell = docx_table._Cell(physical_tcs[col_offset], docx_row.table)  # noqa: SLF001
        _pack_table_cell(
            docx_doc,
            docx_cell,
            cell,  # ty:ignore[invalid-argument-type] Cell already awaited.
            default_comment_author,
        )
        col_offset += cell.grid_span or 1  # ty:ignore[unresolved-attribute] already awaited.


def _pack_table_cell(
    docx_doc: docx_document.Document,
    docx_cell: docx_table._Cell,
    cell: table.TableCell,
    default_comment_author: str | None,
) -> None:
    """Pack a TableCell.

    Args:
        docx_doc: The python-docx Document (needed for comment API).
        docx_cell: The python-docx table cell.
        cell: The declarative TableCell.
        default_comment_author: Default author for comments.
    """
    if cell.children:
        for idx, child in enumerate(cell.children):  # ty:ignore[invalid-argument-type] callables have been resolved.
            if idx == 0 and isinstance(child, paragraph.Paragraph):
                _pack_paragraph_into_existing(
                    docx_doc, docx_cell.paragraphs[0], child, default_comment_author
                )
            else:
                _pack_block_element(docx_doc, docx_cell, child, default_comment_author)  # ty:ignore[invalid-argument-type] already awaited.

    if cell.grid_span is not None and cell.grid_span > 1:
        tc = docx_cell._tc  # noqa: SLF001
        tc.grid_span = cell.grid_span
        tr = tc.getparent()
        tcs = tr.findall(qn("w:tc"))
        tc_index = tcs.index(tc)
        for surplus_tc in tcs[tc_index + 1 : tc_index + cell.grid_span]:
            tr.remove(surplus_tc)

    if cell.vmerge is not None:
        docx_cell._tc.vMerge = (  # noqa: SLF001
            docx_simpletypes.ST_Merge.RESTART
            if cell.vmerge == "restart"
            else docx_simpletypes.ST_Merge.CONTINUE
        )
