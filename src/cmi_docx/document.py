"""Extends a python-docx Word document with additional functionality."""

import pathlib

from docx import document
from docx.text import paragraph as docx_paragraph

from cmi_docx import paragraph, run


class ExtendDocument:
    """Extends a python-docx Word document with additional functionality."""

    def __init__(self, document: document.Document) -> None:
        """Initializes a DocxSearch object for finding text."""
        self.document = document

    def find_in_paragraphs(self, needle: str) -> list[paragraph.FindParagraph]:
        """Finds the indices of a text relative to the paragraphs.

        Args:
            needle: The text to find.

        Returns:
            The indices of the text in the document.
        """
        return [
            paragraph.ExtendParagraph(para).find_in_paragraph(needle)
            for para in self.all_paragraphs
        ]

    def find_in_runs(self, needle: str) -> list[run.FindRun]:
        """Finds the indices of a text relative to the document's runs.

        Args:
            needle: The text to find.

        Returns:
            The indices of the text in the document.
        """
        run_finds: list[run.FindRun] = []
        for document_paragraph in self.all_paragraphs:
            run_finds.extend(
                paragraph.ExtendParagraph(document_paragraph).find_in_runs(needle)
            )
        return run_finds

    def replace(self, needle: str, replace: str) -> None:
        """Finds and replaces text in a Word document.

        Args:
            needle: The text to find.
            replace: The text to replace.

        """
        run_finder = self.find_in_runs(needle)
        run_finder.sort(
            key=lambda x: (x.run_indices[0], x.character_indices[0]), reverse=True
        )

        for run_find in run_finder:
            run_find.replace(replace)

    def insert_paragraph(
        self,
        paragraph: docx_paragraph.Paragraph,
        index: int,
    ) -> None:
        """Inserts a paragraph at a given index.

        Args:
            paragraph: The paragraph to insert.
            index: The index to insert the paragraph at.
        """
        new_paragraph = self._insert_empty_paragraph(index)
        for paragraph_run in paragraph.runs:
            new_paragraph.add_run(paragraph_run.text, paragraph_run.style)

    def insert_image(
        self,
        index: int,
        image_path: str | pathlib.Path,
        width: int | None = None,
        height: int | None = None,
    ) -> None:
        """Inserts an image at a given paragraph index.

        Args:
            index: The paragraph index to insert the image at.
            image_path: The path to the image to insert.
            width: The width of the image.
            height: The height of the image.
        """
        new_paragraph = self._insert_empty_paragraph(index)
        run = new_paragraph.add_run()
        run.add_picture(str(image_path), width=width, height=height)
        return new_paragraph

    @property
    def all_paragraphs(self) -> list[docx_paragraph.Paragraph]:
        """Returns all paragraphs in the document, including headers and footers."""
        all_paragraphs = list(self.document.paragraphs)

        for section in self.document.sections:
            all_paragraphs.extend(
                (*section.footer.paragraphs, *section.header.paragraphs)
            )
        return all_paragraphs

    def _insert_empty_paragraph(self, index: int) -> docx_paragraph.Paragraph:
        """Inserts an empty paragraph at a given index.

        Args:
            index: The index to insert the paragraph at.

        Returns:
            The new paragraph.
        """
        n_paragraphs = len(self.document.paragraphs)
        if index > n_paragraphs:
            raise ValueError(f"Index {index} is out of range.")

        if index == n_paragraphs:
            new_paragraph = self.document.add_paragraph()
        else:
            new_paragraph = new_paragraph = self.document.paragraphs[
                index
            ]._insert_paragraph_before()

        return new_paragraph